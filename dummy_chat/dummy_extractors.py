import requests
import pdfplumber
from PyPDF2 import PdfReader
from docx import Document
from bs4 import BeautifulSoup
import pytesseract
from typing import Optional, Dict, Tuple, List
from openai import OpenAI
import logging
from urllib.parse import urlparse
import time
import certifi
import os
from dataclasses import dataclass
from enum import Enum

logger = logging.getLogger(__name__)

OCR_MIN_TEXT_LEN = 200

# === CONFIGURATION ===

MONGODB_URI = os.getenv("MONGODB_URI")

DB_NAME = os.getenv("KB_DB_NAME", "skylix_rag")
COLLECTION_NAME = os.getenv("KB_COLLECTION_NAME", "kb_chunks")
VECTOR_INDEX_NAME = os.getenv("KB_VECTOR_INDEX", "kb_chunks_vector")
STATUS_FILTER = os.getenv("KB_STATUS_FILTER", "published")

# Lazy initialization - runs only when called
_openai_client = None


def get_openai_client():
    global _openai_client
    if _openai_client is None:
        _openai_client = OpenAI(
            api_key=os.getenv("OPENAI_API_KEY"))  # ✅ Works!
    return _openai_client


@dataclass
class ExtractionConfig:
    """Centralized configuration for extraction behavior"""
    # Request settings
    max_retries: int = 3
    timeout: int = 15
    user_agent: str = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"

    # Content validation
    min_content_length: int = 100
    min_paragraph_length: int = 30

    # Extraction strategies
    try_newspaper: bool = True
    try_trafilatura: bool = True
    try_readability: bool = True
    try_playwright: bool = False  # Set to True if you install playwright
    try_jina_reader: bool = False  # Set to True if you have API key

    # Jina Reader API (if using)
    jina_api_key: Optional[str] = None
    jina_api_url: str = "https://r.jina.ai/"

    # Playwright settings
    playwright_timeout: int = 30000  # 30 seconds
    playwright_wait_until: str = "networkidle"  # or "load", "domcontentloaded"

    # Content selectors (priority order)
    content_selectors: List[str] = None

    # Blocklist patterns for navigation/UI elements
    ui_text_blocklist: List[str] = None

    def __post_init__(self):
        if self.content_selectors is None:
            self.content_selectors = [
                'article',
                'main',
                '[role="main"]',
                '.article-content',
                '.post-content',
                '.entry-content',
                '.content-body',
                '.content-main',
                '#content',
                '.article-body',
                '.post-body',
                '.main-content',
                '.page-content',
                '.entry',
                '.post',
            ]

        if self.ui_text_blocklist is None:
            self.ui_text_blocklist = [
                'share', 'tweet', 'email', 'print', 'menu', 'search',
                'login', 'sign up', 'subscribe', 'newsletter', 'follow us',
                'read more', 'click here', 'learn more', 'advertisement',
                'sponsored', 'related articles', 'recommended', 'trending',
                'popular posts', 'categories', 'tags', 'comments', 'leave a comment'
            ]

        # Load from environment if available
        if not self.jina_api_key:
            self.jina_api_key = os.getenv("JINA_READER_API_KEY")


class ExtractionMethod(Enum):
    """Track which method successfully extracted content"""
    NEWSPAPER = "newspaper3k"
    TRAFILATURA = "trafilatura"
    READABILITY = "readability"
    PLAYWRIGHT = "playwright"
    JINA_READER = "jina_reader"
    BEAUTIFULSOUP = "beautifulsoup"
    FALLBACK = "fallback"


# Global config instance
CONFIG = ExtractionConfig()


# === KEEP YOUR ORIGINAL EXTRACTORS UNCHANGED ===

def extract_pdf(file) -> str:
    text = ""

    # pdfplumber (tables + layout)
    try:
        file.file.seek(0)
        with pdfplumber.open(file.file) as pdf:
            for page in pdf.pages:
                text += (page.extract_text() or "") + "\n"
    except Exception:
        pass

    # PyPDF2 fallback
    if len(text.strip()) < OCR_MIN_TEXT_LEN:
        try:
            file.file.seek(0)
            reader = PdfReader(file.file)
            for page in reader.pages:
                text += page.extract_text() or ""
        except Exception:
            pass

    # OCR fallback
    if len(text.strip()) < OCR_MIN_TEXT_LEN:
        try:
            file.file.seek(0)
            with pdfplumber.open(file.file) as pdf:
                for page in pdf.pages:
                    image = page.to_image(resolution=300).original
                    text += pytesseract.image_to_string(image)
        except Exception:
            pass

    return text.strip()


def extract_docx(file) -> str:
    file.file.seek(0)
    doc = Document(file.file)
    return "\n".join(p.text for p in doc.paragraphs)


def extract_txt(file) -> str:
    file.file.seek(0)
    return file.file.read().decode("utf-8", errors="ignore")


# === ENHANCED LINK EXTRACTION WITH MULTIPLE LIBRARIES ===

class LinkExtractionError(Exception):
    """Custom exception for link extraction failures"""
    pass


@dataclass
class ExtractionResult:
    """Structured result from extraction"""
    content: str
    method: ExtractionMethod
    metadata: Dict[str, Optional[str]]
    success: bool = True
    error: Optional[str] = None


def _validate_url(url: str) -> bool:
    """Validate URL format"""
    try:
        result = urlparse(url)
        return all([result.scheme in ['http', 'https'], result.netloc])
    except Exception:
        return False


def _get_request_headers(config: ExtractionConfig) -> Dict[str, str]:
    """Generate request headers dynamically"""
    return {
        'User-Agent': config.user_agent,
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
        'Accept-Language': 'en-US,en;q=0.9',
        'Accept-Encoding': 'gzip, deflate',
        'Connection': 'keep-alive',
        'Upgrade-Insecure-Requests': '1',
    }


def _fetch_with_retry(
    url: str,
    config: ExtractionConfig
) -> Tuple[str, str]:
    """
    Fetch URL with retry logic and multiple SSL strategies
    Returns: (html_content, content_type)
    """
    headers = _get_request_headers(config)

    # Strategy 1: Standard request with certifi
    for attempt in range(config.max_retries):
        try:
            logger.info(
                f"Attempt {attempt + 1}/{config.max_retries}: Fetching {url}")
            response = requests.get(
                url,
                headers=headers,
                timeout=config.timeout,
                verify=certifi.where(),
                allow_redirects=True
            )
            response.raise_for_status()
            content_type = response.headers.get('content-type', '').lower()
            logger.info(
                f"Successfully fetched {url} - Content-Type: {content_type}")
            return response.text, content_type
        except requests.exceptions.SSLError as e:
            logger.warning(
                f"SSL error on attempt {attempt + 1}: {str(e)[:100]}")
            if attempt < config.max_retries - 1:
                time.sleep(2 ** attempt)  # Exponential backoff
            continue
        except requests.exceptions.Timeout as e:
            logger.warning(f"Timeout on attempt {attempt + 1}: {e}")
            if attempt < config.max_retries - 1:
                time.sleep(2 ** attempt)
            continue
        except requests.exceptions.RequestException as e:
            logger.warning(
                f"Request error on attempt {attempt + 1}: {str(e)[:100]}")
            if attempt < config.max_retries - 1:
                time.sleep(2 ** attempt)
            continue

    # Strategy 2: Fallback without SSL verification
    try:
        logger.warning(
            f"All verified attempts failed. Trying without SSL verification for {url}")
        import urllib3
        urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

        response = requests.get(
            url,
            headers=headers,
            timeout=config.timeout,
            verify=False,
            allow_redirects=True
        )
        response.raise_for_status()
        content_type = response.headers.get('content-type', '').lower()
        return response.text, content_type
    except Exception as e:
        raise LinkExtractionError(
            f"Failed to fetch URL after all retries: {str(e)}")


def _extract_with_newspaper(url: str, html: str, config: ExtractionConfig) -> Optional[ExtractionResult]:
    """Extract using newspaper3k library"""
    if not config.try_newspaper:
        return None

    try:
        from newspaper import Article

        logger.info("Attempting extraction with newspaper3k")
        article = Article(url)
        article.set_html(html)
        article.parse()

        content = article.text

        if len(content.strip()) >= config.min_content_length:
            metadata = {
                'title': article.title,
                'authors': ', '.join(article.authors) if article.authors else None,
                'publish_date': str(article.publish_date) if article.publish_date else None,
                'top_image': article.top_image,
                'source_domain': urlparse(url).netloc
            }
            logger.info(
                f"newspaper3k successfully extracted {len(content)} characters")
            return ExtractionResult(
                content=content,
                method=ExtractionMethod.NEWSPAPER,
                metadata=metadata
            )
    except ImportError:
        logger.debug("newspaper3k not installed, skipping")
    except Exception as e:
        logger.debug(f"newspaper3k extraction failed: {e}")

    return None


def _extract_with_trafilatura(url: str, html: str, config: ExtractionConfig) -> Optional[ExtractionResult]:
    """Extract using trafilatura library"""
    if not config.try_trafilatura:
        return None

    try:
        import trafilatura

        logger.info("Attempting extraction with trafilatura")
        content = trafilatura.extract(
            html,
            include_comments=False,
            include_tables=True,
            no_fallback=False
        )

        if content and len(content.strip()) >= config.min_content_length:
            # Extract metadata using trafilatura
            metadata_obj = trafilatura.extract_metadata(html)
            metadata = {
                'title': metadata_obj.title if metadata_obj else None,
                'author': metadata_obj.author if metadata_obj else None,
                'publish_date': metadata_obj.date if metadata_obj else None,
                'source_domain': urlparse(url).netloc
            }
            logger.info(
                f"trafilatura successfully extracted {len(content)} characters")
            return ExtractionResult(
                content=content,
                method=ExtractionMethod.TRAFILATURA,
                metadata=metadata
            )
    except ImportError:
        logger.debug("trafilatura not installed, skipping")
    except Exception as e:
        logger.debug(f"trafilatura extraction failed: {e}")

    return None


def _extract_with_readability(url: str, html: str, config: ExtractionConfig) -> Optional[ExtractionResult]:
    """Extract using readability-lxml library"""
    if not config.try_readability:
        return None

    try:
        from readability import Document as ReadabilityDocument

        logger.info("Attempting extraction with readability-lxml")
        doc = ReadabilityDocument(html)

        # Get cleaned HTML and convert to text
        clean_html = doc.summary()
        soup = BeautifulSoup(clean_html, 'html.parser')
        content = soup.get_text(separator='\n', strip=True)

        if len(content.strip()) >= config.min_content_length:
            metadata = {
                'title': doc.title(),
                'source_domain': urlparse(url).netloc
            }
            logger.info(
                f"readability-lxml successfully extracted {len(content)} characters")
            return ExtractionResult(
                content=content,
                method=ExtractionMethod.READABILITY,
                metadata=metadata
            )
    except ImportError:
        logger.debug("readability-lxml not installed, skipping")
    except Exception as e:
        logger.debug(f"readability-lxml extraction failed: {e}")

    return None


def _extract_with_playwright(url: str, config: ExtractionConfig) -> Optional[ExtractionResult]:
    """Extract using Playwright for JavaScript-rendered content"""
    if not config.try_playwright:
        return None

    try:
        from playwright.sync_api import sync_playwright

        logger.info(
            "Attempting extraction with Playwright (JavaScript rendering)")

        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page()

            # Navigate and wait for content
            page.goto(url, wait_until=config.playwright_wait_until,
                      timeout=config.playwright_timeout)

            # Wait a bit for any dynamic content
            page.wait_for_timeout(2000)

            # Get the fully rendered HTML
            html = page.content()

            # Extract text from main content
            content = page.evaluate("""
                () => {
                    const main = document.querySelector('article, main, [role="main"], .article-content, .post-content');
                    return main ? main.innerText : document.body.innerText;
                }
            """)

            # Get title
            title = page.title()

            browser.close()

            if content and len(content.strip()) >= config.min_content_length:
                metadata = {
                    'title': title,
                    'source_domain': urlparse(url).netloc
                }
                logger.info(
                    f"Playwright successfully extracted {len(content)} characters")
                return ExtractionResult(
                    content=content,
                    method=ExtractionMethod.PLAYWRIGHT,
                    metadata=metadata
                )
    except ImportError:
        logger.debug("Playwright not installed, skipping")
    except Exception as e:
        logger.debug(f"Playwright extraction failed: {e}")

    return None


def _extract_with_jina_reader(url: str, config: ExtractionConfig) -> Optional[ExtractionResult]:
    """Extract using Jina Reader API"""
    if not config.try_jina_reader or not config.jina_api_key:
        return None

    try:
        logger.info("Attempting extraction with Jina Reader API")

        headers = {
            'Authorization': f'Bearer {config.jina_api_key}',
            'X-Return-Format': 'text'
        }

        response = requests.get(
            f"{config.jina_api_url}{url}",
            headers=headers,
            timeout=config.timeout
        )
        response.raise_for_status()

        content = response.text

        if len(content.strip()) >= config.min_content_length:
            # Parse title from content if available (Jina often includes it)
            lines = content.split('\n')
            title = lines[0] if lines else None

            metadata = {
                'title': title,
                'source_domain': urlparse(url).netloc
            }
            logger.info(
                f"Jina Reader successfully extracted {len(content)} characters")
            return ExtractionResult(
                content=content,
                method=ExtractionMethod.JINA_READER,
                metadata=metadata
            )
    except Exception as e:
        logger.debug(f"Jina Reader extraction failed: {e}")

    return None


def _extract_with_beautifulsoup(html: str, url: str, config: ExtractionConfig) -> ExtractionResult:
    """Extract using BeautifulSoup with smart content detection"""
    logger.info("Attempting extraction with BeautifulSoup")

    soup = BeautifulSoup(html, 'html.parser')

    # Remove unwanted elements
    for tag in soup(['script', 'style', 'noscript', 'iframe', 'svg', 'img']):
        tag.decompose()

    text_parts = []

    # Try content selectors in priority order
    for selector in config.content_selectors:
        elements = soup.select(selector)
        if elements:
            logger.info(f"Found content using selector: {selector}")
            for element in elements:
                # Remove nested unwanted elements
                for tag in element.find_all(['nav', 'footer', 'header', 'aside', 'form']):
                    tag.decompose()

                text = element.get_text(separator='\n', strip=True)
                if len(text) > config.min_content_length:
                    text_parts.append(text)

            if text_parts:
                break

    # Fallback: Extract paragraphs
    if not text_parts:
        logger.info("Selector-based extraction failed, extracting paragraphs")
        paragraphs = soup.find_all(
            ['p', 'h1', 'h2', 'h3', 'h4', 'li', 'blockquote'])
        for p in paragraphs:
            text = p.get_text(strip=True)
            if len(text) > config.min_paragraph_length:
                text_parts.append(text)

    content = '\n\n'.join(text_parts)

    # Extract metadata
    metadata = _extract_metadata(soup, url)

    return ExtractionResult(
        content=content,
        method=ExtractionMethod.BEAUTIFULSOUP,
        metadata=metadata
    )


def _extract_metadata(soup: BeautifulSoup, url: str) -> Dict[str, Optional[str]]:
    """Extract metadata from HTML"""
    metadata = {
        'title': None,
        'description': None,
        'author': None,
        'publish_date': None,
        'source_domain': urlparse(url).netloc
    }

    # Title
    title_tag = soup.find('title')
    if title_tag:
        metadata['title'] = title_tag.get_text(strip=True)

    # Open Graph title (often better than <title>)
    og_title = soup.find('meta', property='og:title')
    if og_title and og_title.get('content'):
        metadata['title'] = og_title['content'].strip()

    # Description
    meta_desc = soup.find('meta', attrs={'name': 'description'}) or \
        soup.find('meta', property='og:description')
    if meta_desc and meta_desc.get('content'):
        metadata['description'] = meta_desc['content'].strip()

    # Author
    meta_author = soup.find('meta', attrs={'name': 'author'}) or \
        soup.find('meta', property='article:author')
    if meta_author and meta_author.get('content'):
        metadata['author'] = meta_author['content'].strip()

    # Date
    meta_date = soup.find('meta', property='article:published_time') or \
        soup.find('meta', attrs={'name': 'publishdate'}) or \
        soup.find('time')
    if meta_date:
        if meta_date.get('content'):
            metadata['publish_date'] = meta_date['content'].strip()
        elif meta_date.get('datetime'):
            metadata['publish_date'] = meta_date['datetime'].strip()

    return metadata


def _clean_extracted_content(text: str, config: ExtractionConfig) -> str:
    """Clean extracted content intelligently"""
    lines = text.split('\n')
    cleaned_lines = []
    prev_line = ""

    for line in lines:
        line = line.strip()

        # Skip empty lines
        if not line:
            continue

        # Skip very short lines
        if len(line) < 10:
            continue

        # Skip duplicate consecutive lines
        if line == prev_line:
            continue

        # Skip UI/navigation elements (case-insensitive)
        if any(blocked.lower() in line.lower() for blocked in config.ui_text_blocklist):
            if len(line) < 50:  # Only skip short lines with blocked text
                continue

        cleaned_lines.append(line)
        prev_line = line

    return '\n\n'.join(cleaned_lines)


def extract_link(url: str, custom_config: Optional[ExtractionConfig] = None) -> str:
    """
    Industry-grade link extraction with multiple library fallbacks

    Extraction order:
    1. Jina Reader API (if configured)
    2. Newspaper3k
    3. Trafilatura
    4. Readability
    5. Playwright (if configured)
    6. BeautifulSoup

    Args:
        url: The URL to extract content from
        custom_config: Optional custom configuration

    Returns:
        Extracted and cleaned text content

    Raises:
        LinkExtractionError: If extraction fails after all attempts
    """
    config = custom_config or CONFIG

    # Normalize URL if scheme is missing (e.g., example.com)
    if not _validate_url(url):
        candidate = f"https://{url}"
        if _validate_url(candidate):
            url = candidate
        else:
            raise LinkExtractionError(f"Invalid URL format: {url}")

    logger.info(f"Starting extraction for URL: {url}")
    logger.info(f"Enabled methods: Newspaper={config.try_newspaper}, Trafilatura={config.try_trafilatura}, "
                f"Readability={config.try_readability}, Playwright={config.try_playwright}, "
                f"JinaReader={config.try_jina_reader and bool(config.jina_api_key)}")

    result = None

    # Try Jina Reader first (most reliable, API-based)
    if config.try_jina_reader and config.jina_api_key:
        result = _extract_with_jina_reader(url, config)
        if result and result.content:
            result.content = _clean_extracted_content(result.content, config)
            if len(result.content.strip()) >= config.min_content_length:
                logger.info(
                    f"✅ Successfully extracted using {result.method.value}")
                return _format_final_output(result)

    # Fetch HTML for other methods
    try:
        html, content_type = _fetch_with_retry(url, config)
    except LinkExtractionError:
        raise

    # Check content type
    if 'html' not in content_type:
        if 'pdf' in content_type:
            raise LinkExtractionError(
                "This appears to be a PDF link. Please download and upload as a file instead."
            )
        logger.warning(f"Non-HTML content type: {content_type}")

    # Try Newspaper3k
    if not result:
        result = _extract_with_newspaper(url, html, config)
        if result and len(result.content.strip()) >= config.min_content_length:
            result.content = _clean_extracted_content(result.content, config)
            if len(result.content.strip()) >= config.min_content_length:
                logger.info(
                    f"✅ Successfully extracted using {result.method.value}")
                return _format_final_output(result)

    # Try Trafilatura
    if not result:
        result = _extract_with_trafilatura(url, html, config)
        if result and len(result.content.strip()) >= config.min_content_length:
            result.content = _clean_extracted_content(result.content, config)
            if len(result.content.strip()) >= config.min_content_length:
                logger.info(
                    f"✅ Successfully extracted using {result.method.value}")
                return _format_final_output(result)

    # Try Readability
    if not result:
        result = _extract_with_readability(url, html, config)
        if result and len(result.content.strip()) >= config.min_content_length:
            result.content = _clean_extracted_content(result.content, config)
            if len(result.content.strip()) >= config.min_content_length:
                logger.info(
                    f"✅ Successfully extracted using {result.method.value}")
                return _format_final_output(result)

    # Try Playwright (for JS-heavy sites)
    if config.try_playwright:
        result = _extract_with_playwright(url, config)
        if result and len(result.content.strip()) >= config.min_content_length:
            result.content = _clean_extracted_content(result.content, config)
            if len(result.content.strip()) >= config.min_content_length:
                logger.info(
                    f"✅ Successfully extracted using {result.method.value}")
                return _format_final_output(result)

    # Fallback to BeautifulSoup
    logger.info(
        "All specialized extractors failed or returned insufficient content, using BeautifulSoup")
    result = _extract_with_beautifulsoup(html, url, config)
    result.content = _clean_extracted_content(result.content, config)

    # Final validation
    if len(result.content.strip()) < config.min_content_length:
        raise LinkExtractionError(
            f"Extracted content is too short ({len(result.content)} chars, minimum: {config.min_content_length}). "
            f"The page might be behind a paywall, require JavaScript rendering (try enabling Playwright), "
            f"or contain mostly non-text content. Tried methods: "
            f"{'Jina, ' if config.try_jina_reader else ''}"
            f"{'Newspaper, ' if config.try_newspaper else ''}"
            f"{'Trafilatura, ' if config.try_trafilatura else ''}"
            f"{'Readability, ' if config.try_readability else ''}"
            f"{'Playwright, ' if config.try_playwright else ''}"
            f"BeautifulSoup"
        )

    logger.info(f"✅ Successfully extracted using {result.method.value}")
    return _format_final_output(result)


def _format_final_output(result: ExtractionResult) -> str:
    """Format the final output with metadata"""
    output_parts = []

    # Add title if available
    if result.metadata.get('title'):
        output_parts.append(f"Title: {result.metadata['title']}")

    # Add author if available
    if result.metadata.get('author') or result.metadata.get('authors'):
        author = result.metadata.get(
            'author') or result.metadata.get('authors')
        output_parts.append(f"Author: {author}")

    # Add date if available
    if result.metadata.get('publish_date'):
        output_parts.append(f"Published: {result.metadata['publish_date']}")

    # Add separator if metadata exists
    if output_parts:
        output_parts.append("\n" + "="*50 + "\n")

    # Add content
    output_parts.append(result.content)

    return '\n'.join(output_parts)


# === CONFIGURATION HELPER FUNCTIONS ===

def update_config(**kwargs):
    """
    Update global configuration

    Example:
        update_config(try_playwright=True, min_content_length=50)
    """
    global CONFIG
    for key, value in kwargs.items():
        if hasattr(CONFIG, key):
            setattr(CONFIG, key, value)
            logger.info(f"Updated config: {key} = {value}")
        else:
            logger.warning(f"Unknown config key: {key}")


def get_config() -> ExtractionConfig:
    """Get current configuration"""
    return CONFIG


def reset_config():
    """Reset configuration to defaults"""
    global CONFIG
    CONFIG = ExtractionConfig()
    logger.info("Configuration reset to defaults")
